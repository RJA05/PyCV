from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
import os
from docx2pdf import convert as convert_to_pdf

#========================================= DOCX edit =========================================
#### Open new doc ###
doc=Document()

### save the generated CV to specified or unspecified name ###
def save_document(profile_name:str, filename="") -> str:
    invalid_chars = '<>:"/\\|?*'
    clean_filename = ''.join(c for c in filename if c not in invalid_chars)

    if clean_filename=="":
        counter=1
        while os.path.exists(str(profile_name) +"-"+ str(counter) +".docx"):
            counter+=1
        try:
            doc.save(str(profile_name)+"-" + str(counter) +".docx")
            return str(profile_name)+"-" + str(counter) +".docx"
        except:
            print("ERROR: could not save file")
    else:
        try:
            doc.save(str(clean_filename)+".docx")
            return str(clean_filename)+".docx"
        except:
            print("ERROR: could not save file")

### set doc margins and other parameters maybe idk yet using section 0 ###
def doc_style(top: float=0.6,bottom: float=0.5,left: float=0.6,right: float=0.6,space_after:float=2,space_before:float=2,line_spacing: int=12):
    section = doc.sections[0]
    ### margins ###
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    ### spacing ###
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing =Pt(line_spacing)
    ### font ###
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

### line called in each section, pass the paragraph name to use ###
def section_line(paragraph,thickness: int=8):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    borders.append(bottom)
    pPr.append(borders)

### adds hyperlink text ###
def add_hyperlink(paragraph, text:str, url:str,f_size:int):
    part = paragraph.part
    r_id = part.relate_to(url,"http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '#6e9cc4')
    rPr.append(color)
    # Set font size (in half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(f_size * 2)))
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

### write the name and contact info to doc ###
### contact_info={"email":"" , "phone":"", "linkedin": "", "webpage":"", "address":""} ###
def contact_section(name:str, contact_info: dict={"email":"" , "phone":"", "linkedin": "", "webpage":"", "address":""}):
    ### name ###
    capName=""
    for i in name.split():
        capName=capName + i.capitalize() +" "
    n=doc.add_paragraph()
    run=n.add_run(capName)
    run.bold = True
    run.font.size = Pt(12)
    n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ### contact ingo ###
    if contact_info=={}:
        return
    else:
        ks=list(contact_info.keys())
        key_num=len(ks)
        c1=doc.add_paragraph()
        for i in range(key_num):
            if i>2:
                break
            if ks[i]== "linkedin":
                run=c1.add_run("Linkedin: ")
                run.font.size = Pt(12)
                add_hyperlink(c1,"/"+str(contact_info[ks[i]]),"https://www.linkedin.com/in/"+str(contact_info[ks[i]]),12)
            elif ks[i]== "website":
                add_hyperlink(c1,str(contact_info[ks[i]]),str(contact_info[ks[i]]),12)
            else:
                run=c1.add_run(str(contact_info[ks[i]]))
                run.font.size = Pt(12)
            if i+1<3 and i<key_num-1:
                run=c1.add_run("   |   ")
                run.font.size = Pt(12)
            c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if key_num >3:
            c2=doc.add_paragraph()
            for i in range(3,key_num,1):
                if ks[i]== "linkedin":
                    run=c2.add_run("Linkedin: ")
                    run.font.size = Pt(12)
                    add_hyperlink(c2,"/"+str(contact_info[ks[i]]),"https://www.linkedin.com/in/"+str(contact_info[ks[i]]),12)
                elif ks[i]== "webpage":
                    add_hyperlink(c2,str(contact_info[ks[i]]),str(contact_info[ks[i]]),12)
                else:
                    run=c2.add_run(str(contact_info[ks[i]]))
                    run.font.size = Pt(12)
                if i<key_num-1:
                    run=c2.add_run("   |   ")
                    run.font.size = Pt(12)
            c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run=c2.add_run()
        run.add_break() 

### write Professional summary section
def text_section(text:str, text_type=None, title=""):
    match text_type:
        case "summary":
            s=doc.add_paragraph()
            run=s.add_run("Professional summary: ")
            run.font.size = Pt(10)
            run.bold = True
            Brun=s.add_run(str(text.capitalize()))
            Brun.font.size = Pt(10)
            Brun.add_break() 
            s.paragraph_format.left_indent = Inches(0.5)
            s.paragraph_format.right_indent = Inches(0.5)
            s.paragraph_format.first_line_indent = Inches(0.5)
        case _:
            t=doc.add_paragraph()
            run=t.add_run(title.capitalize())
            run.bold = True
            run.font.size = Pt(12)
            section_line(t)
            s=doc.add_paragraph()
            Brun=s.add_run(str(text.capitalize()))
            Brun.add_break() 

### adds the education section, can list various degrees or titles ###
def education_section(education={"degree":["time","school"]}):
    t=doc.add_paragraph()
    run=t.add_run("Education")
    run.bold = True
    run.font.size = Pt(12)
    section_line(t)
    titles=list(education.keys())
    for title in titles:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        row = table.rows[0]
        cell_title = row.cells[0]
        cell_time = row.cells[1]
        run_title = cell_title.paragraphs[0].add_run(title.capitalize())
        run_title.bold = True
        run_time = cell_time.paragraphs[0].add_run(education[title][0])
        cell_time.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        s = doc.add_paragraph(education[title][1].capitalize())
        s.paragraph_format.left_indent = Inches(0.25)
    s.add_run().add_break()

### writes the skills section and all skills ###
def list_section(items: list[str], title:str, list_type:str=None, item_names:list[str]=[""]):
    match list_type:
        case "bullets":
            s=doc.add_paragraph()
            run=s.add_run(title.capitalize())
            run.bold = True
            run.font.size = Pt(12)
            section_line(s)
            for item in items:
                sk=doc.add_paragraph(item.capitalize(),style="List Bullet")
                sk.paragraph_format.left_indent = Inches(0.5)
            s=sk.add_run()
            s.add_break()
        case "numbers":
            s=doc.add_paragraph()
            run=s.add_run(title.capitalize())
            run.bold = True
            run.font.size = Pt(12)
            section_line(s)
            for item in items:
                sk=doc.add_paragraph(item.capitalize(),style="List Number")
                sk.paragraph_format.left_indent = Inches(0.5)
            s=sk.add_run()
            s.add_break()
        case "names":
            s=doc.add_paragraph()
            run=s.add_run(title.capitalize())
            run.bold = True
            run.font.size = Pt(12)
            section_line(s)
            if len(item_names) != len(items):
                for i in range(abs(len(item_names) - len(items))):
                    if len(item_names) < len(items):
                        item_names.append("undefined")
                    else:
                        items.append("undefined")

            for i in range(len(items)):
                l=doc.add_paragraph()
                l.paragraph_format.left_indent = Inches(0.25)
                t=l.add_run(item_names[i].capitalize()+": ")
                t.bold=True
                it=l.add_run(items[i].capitalize())
            it.add_break()
        case _:
            s=doc.add_paragraph()
            run=s.add_run(title.capitalize())
            run.bold = True
            run.font.size = Pt(12)
            section_line(s)
            for item in items:
                sk=doc.add_paragraph(item.capitalize())
                sk.paragraph_format.left_indent = Inches(0.25)
            s=sk.add_run()
            s.add_break()

def work_section(section_title="Work", projects={"name":["time",["skill 1", "skill 2"]]}):
    s=doc.add_paragraph()
    run=s.add_run(section_title.capitalize())
    run.bold = True
    run.font.size = Pt(12)
    section_line(s)
    names=list(projects.keys())
    for project in names:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        row = table.rows[0]
        cell_title = row.cells[0]
        cell_time = row.cells[1]
        p=cell_title.paragraphs[0].add_run(project.capitalize())
        p.bold=True
        t=cell_time.paragraphs[0].add_run(projects[project][0])
        cell_time.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        t.add_break
        for skill in projects[project][1]:
            s=doc.add_paragraph(skill.capitalize(), style="List Bullet")
            s.paragraph_format.left_indent = Inches(0.5)
    s.add_run().add_break()

#========================================= Profiles =========================================
def find_profiles()-> list:
    profiles:list[str]=[]
    with os.scandir() as entries:
        for entry in entries:
            if entry.is_file() and ".pyCV-profile" in entry.name:  # Check if it's a file
                profiles.append(entry.name)
    return profiles

def read_profile(filename:str)-> dict:#{"title": [ type,["content line 1", "content line 2", ...],{"set 1": [1,2,3], "set 2":"bla bla bla", ...}] }
    # the dict of the entire return has the following structure {"title": [ type,["content line 1", "content line 2", ...],{"set 1": [1,2,3], "set 2":"bla bla bla", ...}] }
    section_type:str=None
    section_title:str=None
    read_state:str=None
    profile:dict={}
    start_set:bool=False
    set_num=1
    with open(filename,"rt") as data_file:
        for line in data_file:
            if "$$" in line:
                read_state="title"
                section_type=line[3:-1]
            elif "title:" in line and read_state=="title":
                section_title=line[6:-1]
                profile[section_title]=[section_type]
                read_state="info"
            elif line=="\n":
                read_state=None
            elif "##" in line:
                read_state="sets"
                profile[section_title].append({})
            elif read_state=="sets":
                if "set:" in line:
                    set_num=line[-2]
                    profile[section_title][-1][set_num]=[]
                elif "{" in line:
                    start_set=True
                elif"}"in line:
                    start_set=False
                else:
                    profile[section_title][-1][set_num].append(line[:-1])
            elif read_state=="info" and "##"not in line:
                if ("title" not in line) and ("## SETS" not in line):
                    if len(profile[section_title])<2:
                        profile[section_title].append({})
                    data=line.split(":")
                    profile[section_title][1][data[0]]=data[1][:-1]
    return profile

#========================================= Profile write to doc =========================================
def write_doc(profile:dict, set_selections:dict, filename:str=""):
    doc_style()
    sections=list(profile.keys())
    for section in sections:
        section_type=profile[section][0]
        match section_type:
            case "NAME":
                name=profile[section][1]["name"]
            case "CONTACT":
                contact_info=profile[section][1]
                contact_section(name, contact_info)
            case "TEXT":
                text_type = profile[section][1]["type"]
                selection=set_selections[section]
                if text_type=='summary':
                    text_section(profile[section][2][selection][0],"summary")
                else:
                    text_section(profile[section][2][selection][0],"plain",str(section))
            case "WORK":
                projects={}
                selection=set_selections[section]
                project_read=profile[section][1][selection]
                for project in project_read:
                    project_info= project.split(":")
                    projects[project_info[0]]=[project_info[1],project_info[2].split(",")]
                work_section(str(section),projects)
            case "EDUCATION":
                titles={}
                title_read=profile[section][1]["1"]
                for title in title_read:
                    title_info=title.split(":")
                    titles[title_info[0]]=[title_info[1],title_info[2]]
                education_section(titles)
            case "LIST":
                list_type=profile[section][1]["type"]
                selection=set_selections[section]
                items=profile[section][2][selection]
                print(list_type)
                match list_type:
                    case "bullets":
                        list_section(items,str(section),"bullets")
                    case "numbers":
                        list_section(items,str(section),"numbers")
                    case "names":
                        item_list=[]
                        item_name_list=[]
                        for item in items:
                            split_item= item.split(":")
                            item_name_list.append(split_item[0])
                            item_list.append(split_item[1])
                        list_section(item_list,str(section),"names",item_name_list)
                    case "plain":
                        list_section(items,str(section))


    file_saved=save_document(name,filename)
    print(f"saved to {file_saved}")


if __name__ == '__main__': ### for testing
    # print("test start")
    # doc_style()
    # contact_section("glerb sgurp",{"email":"email1111" , "phone":"+52 (123)45678", "linkedin": "cesarbermudez2003", "webpage":"webpage.com", "address":"address"})
    # text_section("summary summary summary summary summary summary summary summary summary summary summary summary summary summary summary summary summary summary summary summary he.","summary")
    # education_section({"bachelors in gurb":["jan 2024 to jan 2021", "tredlurb school"],"masters in gub":["jan 2024 to jan 2021", "tredlurb school"]})
    # list_section(["a","b"],"skills","bullets")
    # list_section(["t","o"],"test","names",["k","p"])
    # list_section(["a","b","c"],"no type")
    # text_section("glibur glibur gliburgliburglibur glibur gliburgliburglibur",title="TEXT :)")
    # work_section()
    # filename=save_document("glorb","glorb")
    #convert_to_pdf(filename)
    profiles=find_profiles()
    example_profile=read_profile(profiles[0])
    selections={"summary":"2", "plain text section":"1", "work section title":"1","list section title":"1","another type of list":"1"}
    write_doc(example_profile,selections,"test")
